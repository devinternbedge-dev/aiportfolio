from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

from pypdf import PdfReader
from docx import Document as DocxDocument

import os

# ==========================================
# Paths
# ==========================================

BASE_DIR = os.path.dirname(
    os.path.dirname(
        os.path.abspath(__file__)
    )
)

DATA_PATH = os.path.join(
    BASE_DIR,
    "data"
)

VECTORSTORE_PATH = os.path.join(
    BASE_DIR,
    "vectorstore"
)

documents = []

# ==========================================
# Load TXT / PDF / DOCX Files
# ==========================================

for file in os.listdir(DATA_PATH):

    file_path = os.path.join(
        DATA_PATH,
        file
    )

    print(f"Loading: {file}")

    try:

        # ==========================
        # TXT
        # ==========================

        if file.endswith(".txt"):

            with open(
                file_path,
                "r",
                encoding="utf-8"
            ) as f:

                text = f.read()

                if text.strip():

                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": file
                            }
                        )
                    )

        # ==========================
        # PDF
        # ==========================

        elif file.endswith(".pdf"):

            pdf = PdfReader(file_path)

            text = ""

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:

                    text += page_text + "\n"

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file
                        }
                    )
                )

        # ==========================
        # DOCX
        # ==========================

        elif file.endswith(".docx"):

            doc = DocxDocument(file_path)

            text = "\n".join(
                [
                    para.text
                    for para in doc.paragraphs
                ]
            )

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file
                        }
                    )
                )

    except Exception as e:

        print(
            f"Error loading {file}: {str(e)}"
        )

print(
    f"\nDocuments Loaded: {len(documents)}"
)

# ==========================================
# Validate Documents
# ==========================================

if not documents:

    raise ValueError(
        "No valid documents found in data folder."
    )

# ==========================================
# Split Documents
# ==========================================

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=100
)

docs = text_splitter.split_documents(
    documents
)

print(
    f"Chunks Created: {len(docs)}"
)

# ==========================================
# Embeddings
# ==========================================

embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# ==========================================
# Remove Old Vector DB
# ==========================================

if os.path.exists(VECTORSTORE_PATH):

    import shutil

    shutil.rmtree(
        VECTORSTORE_PATH
    )

# ==========================================
# Create FAISS Vector Store
# ==========================================

vectorstore = FAISS.from_documents(
    docs,
    embeddings
)

# ==========================================
# Save Vector Store
# ==========================================

vectorstore.save_local(
    VECTORSTORE_PATH
)

print(
    "\nVector DB Created Successfully"
)

print(
    f"Saved To: {VECTORSTORE_PATH}"
)